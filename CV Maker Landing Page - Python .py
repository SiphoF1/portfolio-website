from docx import Document

def create_cv(name, email, phone, experience):
    doc = Document()
    
    # Add basic information
    doc.add_heading('Curriculum Vitae', level=1)
    doc.add_paragraph(f'Name: {name}')
    doc.add_paragraph(f'Email: {email}')
    doc.add_paragraph(f'Phone: {phone}')
    
    # Add experience
    doc.add_heading('Experience', level=2)
    for exp in experience:
        doc.add_paragraph(exp)
    
    # Add footer
    footer = doc.sections[0].footer
    footer.paragraphs[0].text = "Created by SSI Innovations"
    
    # Save the document
    doc.save('my_cv.docx')
    
    print('CV created successfully!')

# Example data
name = 'John Doe'
email = 'johndoe@example.com'
phone = '123-456-7890'
experience = ['Software Developer at XYZ Company', 'Intern at ABC Inc.']

# Create the CV
create_cv(name, email, phone, experience)